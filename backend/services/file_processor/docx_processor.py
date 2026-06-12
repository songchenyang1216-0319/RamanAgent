from __future__ import annotations

from pathlib import Path

from .base import BaseFileProcessor, ProcessedFile
from .word_utils import convert_legacy_word_to_docx, extract_openxml_text


class DocxFileProcessor(BaseFileProcessor):
    file_type = "docx"
    supported_suffixes = {".docx", ".doc"}

    def process(self, path: Path, *, file_id: str | None = None, **_: object) -> ProcessedFile:
        source_path = path
        conversion_note = None
        if path.suffix.lower() == ".doc":
            source_path, error_message = convert_legacy_word_to_docx(path)
            if source_path is None:
                return self.failure(path, error_message or "旧版 DOC 文件转换失败。")
            conversion_note = f"旧版 DOC 已转换为 {source_path.name} 后解析。"
        try:
            try:
                from docx import Document

                document = Document(str(source_path))
                paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
                table_lines = []
                table_count = len(document.tables)
                for table in document.tables:
                    for row in table.rows:
                        table_lines.append(" | ".join(cell.text.strip() for cell in row.cells))
                text = "\n".join(paragraphs + table_lines)
            except ModuleNotFoundError:
                text, paragraph_count = extract_openxml_text(source_path)
                paragraphs = [line for line in text.splitlines() if line.strip()]
                table_lines = []
                table_count = 0
                if paragraph_count and not paragraphs:
                    paragraphs = [text]
        except Exception as exc:
            return self.failure(path, f"DOCX 文件解析失败：{exc}")
        return ProcessedFile(
            success=True,
            file_type=self.file_type,
            filename=path.name,
            summary=f"已解析 DOCX，提取到 {len(paragraphs)} 个段落、{table_count} 个表格和 {len(table_lines)} 行表格文本。",
            metadata={
                "paragraphs": len(paragraphs),
                "tables": table_count,
                "table_rows": len(table_lines),
                "characters": len(text),
                "suffix": path.suffix.lower(),
                "converted_from_legacy_doc": bool(conversion_note),
                "converted_docx_name": source_path.name if conversion_note else None,
            },
            preview=text[:3000],
            chunks=self.make_chunks(text, file_id=file_id, section="docx"),
            error_message=None,
        )
