from __future__ import annotations

import csv
import html
import json
import os
from pathlib import Path
from typing import Any

from backend.services.file_service import FileCatalogService
from backend.services.file_processor.word_utils import convert_legacy_word_to_docx
from backend.services.workspace_manager import WorkspaceManager
from raman_core.methanol.config import PROJECT_ROOT


class FileConverterService:
    def __init__(self) -> None:
        self.file_catalog = FileCatalogService()
        self.workspace_manager = WorkspaceManager()

    def convert_file(
        self,
        *,
        file_id: str,
        target_format: str,
        user_id: str,
        conversation_id: str,
        is_admin: bool = False,
    ) -> dict[str, Any]:
        target_format = self._normalize_format(target_format)
        source = self.file_catalog.get_file_for_user(file_id, user_id=user_id, is_admin=is_admin)
        if source is None:
            raise FileNotFoundError("文件不存在或无权访问。")
        source_path = (PROJECT_ROOT / str(source.get("path") or "")).resolve()
        self._assert_project_path(source_path)
        if not source_path.exists() or not source_path.is_file():
            raise FileNotFoundError("源文件不存在。")
        suffix = source_path.suffix.lower().lstrip(".")
        actual_format = target_format
        warnings: list[str] = []
        pdf_export_available = True
        content: str | bytes | Path
        if target_format in {"txt", "md", "html"}:
            text = self._extract_text(source_path, suffix)
            content = self._text_to_target(text, target_format, source_path.name)
        elif target_format == "docx":
            content = self._text_to_docx(self._extract_text(source_path, suffix), source_path.name)
        elif target_format == "xlsx":
            if suffix not in {"csv", "tsv"}:
                raise ValueError("当前仅支持 CSV/TSV 转换为 XLSX。")
            content = self._csv_to_xlsx(source_path)
        elif target_format == "csv":
            if suffix not in {"xlsx", "xls"}:
                raise ValueError("当前仅支持 Excel 转换为 CSV。")
            content = self._xlsx_to_csv(source_path)
        elif target_format == "pdf":
            content, actual_format, pdf_export_available, warnings = self._text_to_pdf_or_fallback(self._extract_text(source_path, suffix), source_path.name)
        else:
            raise ValueError(f"暂不支持转换为 {target_format}。")
        converted_name = f"{source_path.stem}.{actual_format}"
        output = self.workspace_manager.save_output_file(user_id, conversation_id, converted_name, content)
        artifact = {
            "artifact_id": output.get("file_id"),
            "type": actual_format,
            "title": f"{source.get('original_filename') or source_path.name} 转换结果",
            "filename": output.get("filename"),
            "mime_type": output.get("mime_type"),
            "url": output.get("download_url"),
            "download_url": output.get("download_url"),
            "preview_url": output.get("preview_url"),
            "created_at": output.get("updated_at"),
        }
        return {
            "success": True,
            "source_file": source,
            "requested_format": target_format,
            "target_format": actual_format,
            "actual_format": actual_format,
            "pdf_export_available": pdf_export_available,
            "warnings": warnings,
            "output_file": output,
            "artifact": artifact,
            "message": "文件转换完成。" if target_format == actual_format else f"PDF 导出不可用，已生成 {actual_format.upper()} fallback。",
        }

    def _normalize_format(self, value: str) -> str:
        fmt = str(value or "").lower().strip().lstrip(".")
        aliases = {"markdown": "md", "text": "txt", "excel": "xlsx", "word": "docx"}
        return aliases.get(fmt, fmt)

    def _assert_project_path(self, path: Path) -> None:
        project_root = PROJECT_ROOT.resolve()
        if path != project_root and project_root not in path.parents:
            raise ValueError("文件路径不合法。")

    def _extract_text(self, path: Path, suffix: str) -> str:
        if suffix in {"txt", "md", "markdown", "log", "csv", "tsv", "json", "yaml", "yml", "py", "js", "ts", "java", "go", "html", "css", "sql", "sh"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix in {"docx", "doc"}:
            if suffix == "doc":
                converted_path, error_message = convert_legacy_word_to_docx(path)
                if converted_path is None:
                    raise ValueError(error_message or "旧版 DOC 文件转换失败。")
                path = converted_path
            try:
                from docx import Document
            except Exception as exc:
                raise ValueError("未安装 python-docx，无法读取 DOCX。") from exc
            doc = Document(str(path))
            parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        if suffix == "pdf":
            try:
                from pypdf import PdfReader
            except Exception as exc:
                raise ValueError("未安装 pypdf，无法读取 PDF。") from exc
            reader = PdfReader(str(path))
            return "\n\n".join((page.extract_text() or "") for page in reader.pages)
        if suffix in {"xlsx", "xls"}:
            return self._xlsx_to_csv(path).read_text(encoding="utf-8", errors="replace")
        raise ValueError(f"当前文件类型 .{suffix} 暂不支持转换。")

    def _text_to_target(self, text: str, target_format: str, source_name: str) -> str:
        if target_format == "txt":
            return text
        if target_format == "md":
            if source_name.lower().endswith((".md", ".markdown")):
                return text
            return f"# {source_name}\n\n```text\n{text}\n```"
        if target_format == "html":
            try:
                import markdown

                return markdown.markdown(text, extensions=["tables", "fenced_code"])
            except Exception:
                return f"<pre>{html.escape(text)}</pre>"
        raise ValueError(f"不支持文本转换目标: {target_format}")

    def _text_to_docx(self, text: str, source_name: str) -> Path:
        try:
            from docx import Document
        except Exception as exc:
            raise ValueError("未安装 python-docx，无法生成 DOCX。") from exc
        target = PROJECT_ROOT / "storage" / "tmp" / f"{Path(source_name).stem}_converted.docx"
        target.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading(Path(source_name).stem, level=1)
        for block in text.split("\n\n"):
            doc.add_paragraph(block.strip())
        doc.save(str(target))
        return target

    def _csv_to_xlsx(self, path: Path) -> Path:
        try:
            from openpyxl import Workbook
        except Exception as exc:
            raise ValueError("未安装 openpyxl，无法生成 XLSX。") from exc
        delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
        target = PROJECT_ROOT / "storage" / "tmp" / f"{path.stem}_converted.xlsx"
        target.parent.mkdir(parents=True, exist_ok=True)
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
            reader = csv.reader(handle, delimiter=delimiter)
            for row in reader:
                ws.append(row)
        wb.save(str(target))
        return target

    def _xlsx_to_csv(self, path: Path) -> Path:
        try:
            from openpyxl import load_workbook
        except Exception as exc:
            raise ValueError("未安装 openpyxl，无法读取 Excel。") from exc
        target = PROJECT_ROOT / "storage" / "tmp" / f"{path.stem}_converted.csv"
        target.parent.mkdir(parents=True, exist_ok=True)
        wb = load_workbook(str(path), read_only=True, data_only=True)
        ws = wb.active
        with target.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(["" if value is None else value for value in row])
        wb.close()
        return target

    def _text_to_pdf_or_fallback(self, text: str, source_name: str) -> tuple[str | Path, str, bool, list[str]]:
        provider = str(os.getenv("PDF_EXPORT_PROVIDER") or "none").strip().lower()
        html_content = self._text_to_target(text, "html", source_name)
        warnings: list[str] = []
        if provider in {"", "none", "disabled", "off", "html"}:
            warnings.append("PDF_EXPORT_PROVIDER 未启用，已生成 HTML fallback。")
            return html_content, "html", False, warnings

        providers = ["weasyprint", "playwright"] if provider == "auto" else [provider]
        last_error = ""
        for item in providers:
            try:
                if item == "weasyprint":
                    return self._html_to_pdf_weasyprint(html_content, source_name), "pdf", True, warnings
                if item == "playwright":
                    return self._html_to_pdf_playwright(html_content, source_name), "pdf", True, warnings
                last_error = f"不支持的 PDF_EXPORT_PROVIDER: {item}"
            except Exception as exc:
                last_error = str(exc)
        warnings.append(f"PDF 导出失败，已生成 HTML fallback。原因：{last_error}")
        return html_content, "html", False, warnings

    def _html_to_pdf_weasyprint(self, html_content: str, source_name: str) -> Path:
        try:
            from weasyprint import HTML  # type: ignore
        except Exception as exc:
            raise ValueError("未安装 weasyprint，无法生成 PDF。") from exc
        target = PROJECT_ROOT / "storage" / "tmp" / f"{Path(source_name).stem}_converted.pdf"
        target.parent.mkdir(parents=True, exist_ok=True)
        HTML(string=html_content).write_pdf(str(target))
        return target

    def _html_to_pdf_playwright(self, html_content: str, source_name: str) -> Path:
        try:
            from playwright.sync_api import sync_playwright  # type: ignore
        except Exception as exc:
            raise ValueError("未安装 playwright，无法生成 PDF。") from exc
        target = PROJECT_ROOT / "storage" / "tmp" / f"{Path(source_name).stem}_converted.pdf"
        temp_html = PROJECT_ROOT / "storage" / "tmp" / f"{Path(source_name).stem}_converted.html"
        target.parent.mkdir(parents=True, exist_ok=True)
        temp_html.write_text(html_content, encoding="utf-8")
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(temp_html.as_uri())
            page.pdf(path=str(target), format="A4", print_background=True)
            browser.close()
        return target
